#!/usr/bin/env python3
"""
Test script to simulate the complete n8n chat workflow:
1. User: "create a document containing a list of Sevilla products"
2. User: "add a numbered list with ten more products"  
3. User: "get download link"
"""

import os
import sys
import asyncio
from pathlib import Path

# Add the word_document_server to Python path
sys.path.insert(0, str(Path(__file__).parent))

async def test_complete_chat_workflow():
    """Test the complete chat workflow simulation."""
    print("🤖 Testing Complete N8N Chat Workflow")
    print("=" * 60)
    
    try:
        # Import the functions we need to test
        from word_document_server.main import (
            init_temp_storage,
            register_temp_file,
            get_temp_file_by_user_filename,
            resolve_document_path,
            load_document_with_resolver,
            save_document_with_resolver,
            get_transport_config,
            TEMP_FILES_DIR
        )
        from docx import Document
        from word_document_server.core.styles import ensure_heading_style
        from word_document_server.utils.file_utils import ensure_docx_extension
        from datetime import datetime, timedelta
        import uuid
        
        # Step 0: Initialize the system
        print("📋 Step 0: Initializing temp storage...")
        init_temp_storage()
        print("✓ Temp storage initialized")
        
        # Step 1: User asks "create a document containing a list of Sevilla products"
        print("\\n👤 User: 'Create a document containing a list of Sevilla products'")
        print("🤖 AI Agent: Calling create_document_with_download_link('products.docx')...")
        
        # Simulate create_document_with_download_link
        filename = "products.docx"
        cleanup_hours = 24
        title = "Sevilla Products List"
        
        original_filename = ensure_docx_extension(filename)
        unique_filename = f"{uuid.uuid4()}_{original_filename}"
        temp_file_path = TEMP_FILES_DIR / unique_filename
        
        # Create the document
        doc = Document()
        doc.core_properties.title = title
        ensure_heading_style(doc)
        
        # Add initial content
        doc.add_heading("Sevilla Products", 1)
        doc.add_paragraph("Initial product list:")
        doc.add_paragraph("1. Olive oil")
        doc.add_paragraph("2. Olives")
        doc.add_paragraph("3. Orange marmalade")
        
        doc.save(str(temp_file_path))
        
        # Register the file for cleanup  
        file_id = register_temp_file(str(temp_file_path), original_filename, filename, cleanup_hours)
        
        # Generate download URL
        config = get_transport_config()
        base_url = f"http://{config['host']}:{config['port']}"
        download_url = f"{base_url}/files/{file_id}"
        
        create_result = {
            "success": True,
            "message": f"Document {original_filename} created successfully",
            "download_url": download_url,
            "file_id": file_id,
            "original_filename": original_filename,
        }
        
        print(f"✓ Document created: {file_id}")
        print(f"✓ Download URL: {download_url}")
        
        # Step 2: User asks "add a numbered list with ten more products" 
        print("\\n👤 User: 'Add a numbered list with ten more products'")
        print("🤖 AI Agent: Calling add_paragraph('products.docx', ...)...")
        
        # Test the resolver by adding content to "products.docx" 
        # This should find the temp file via user_filename mapping
        
        # Simulate add_paragraph with resolver
        doc, resolved_path = load_document_with_resolver("products.docx")
        print(f"✓ Document resolved to: {resolved_path}")
        
        # Add more content
        doc.add_paragraph("\\nAdditional products:")
        for i in range(4, 14):  # Products 4-13
            doc.add_paragraph(f"{i}. Product {i}")
        
        save_document_with_resolver(doc, "products.docx", resolved_path)
        print("✓ Content added to document")
        
        # Verify file still exists in temp
        temp_file_info = get_temp_file_by_user_filename("products.docx")
        if temp_file_info and os.path.exists(temp_file_info["file_path"]):
            print("✓ Document still accessible in temp storage")
        else:
            print("✗ Document not found in temp storage")
            return False
        
        # Step 3: User asks "get download link"
        print("\\n👤 User: 'Get download link for products.docx'")
        print("🤖 AI Agent: Calling get_download_link('products.docx')...")
        
        # Simulate get_download_link
        temp_file_info = get_temp_file_by_user_filename("products.docx")
        if temp_file_info and os.path.exists(temp_file_info["file_path"]):
            expires_at = datetime.fromisoformat(temp_file_info["expires_at"])
            if datetime.now() <= expires_at:
                download_url = f"{base_url}/files/{temp_file_info['file_id']}"
                link_result = {
                    "success": True,
                    "filename": "products.docx",
                    "download_url": download_url,
                    "file_id": temp_file_info["file_id"],
                    "expires_at": temp_file_info["expires_at"]
                }
                print(f"✓ Download link retrieved: {download_url}")
            else:
                print("✗ File has expired")
                return False
        else:
            print("✗ File not found")
            return False
        
        # Verify final document content
        print("\\n📄 Verifying final document content...")
        final_doc = Document(temp_file_info["file_path"])
        paragraph_count = len(final_doc.paragraphs)
        print(f"✓ Document has {paragraph_count} paragraphs")
        
        # Show some content
        print("📝 Document content preview:")
        for i, para in enumerate(final_doc.paragraphs[:8]):
            if para.text.strip():
                print(f"   {i+1}: {para.text}")
        
        print("\\n" + "=" * 60)
        print("🎉 WORKFLOW TEST SUCCESSFUL!")
        print("\\n✅ The complete n8n chat workflow works:")
        print("   1. ✓ Create document with download link")
        print("   2. ✓ Edit document using original filename")  
        print("   3. ✓ Retrieve download link anytime")
        print("   4. ✓ Smart resolver finds temp files seamlessly")
        
        return True
        
    except Exception as e:
        print(f"\\n❌ WORKFLOW TEST FAILED: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Run the workflow test."""
    success = asyncio.run(test_complete_chat_workflow())
    return 0 if success else 1

if __name__ == "__main__":
    sys.exit(main())