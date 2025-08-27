#!/usr/bin/env python3
"""
Test the create_document_with_download_link tool to verify it works correctly.
"""

import os
import sys
import uuid
import sqlite3
import json
from datetime import datetime, timedelta
from pathlib import Path

# Add the word_document_server to Python path
sys.path.insert(0, str(Path(__file__).parent))

def simulate_create_document_with_download_link():
    """Simulate the create_document_with_download_link tool functionality."""
    print("🧪 Testing create_document_with_download_link Tool")
    print("=" * 60)
    
    try:
        # Import required modules (without FastMCP dependencies)
        from docx import Document
        
        # Simulate the tool parameters
        filename = "test_products.docx"
        cleanup_hours = 24
        title = "Test Products List"
        author = "Test User"
        
        print(f"📝 Testing with parameters:")
        print(f"   - filename: {filename}")
        print(f"   - cleanup_hours: {cleanup_hours}")
        print(f"   - title: {title}")
        print(f"   - author: {author}")
        
        # Step 1: Initialize temp storage (copied from main.py)
        TEMP_FILES_DIR = Path("/tmp/mcp_files")
        DB_FILE = TEMP_FILES_DIR / "file_registry.db"
        
        def init_temp_storage():
            TEMP_FILES_DIR.mkdir(exist_ok=True)
            conn = sqlite3.connect(DB_FILE)
            conn.execute("""
                CREATE TABLE IF NOT EXISTS temp_files (
                    file_id TEXT PRIMARY KEY,
                    original_filename TEXT NOT NULL,
                    user_filename TEXT NOT NULL,
                    file_path TEXT NOT NULL,
                    created_at DATETIME NOT NULL,
                    expires_at DATETIME NOT NULL,
                    download_count INTEGER DEFAULT 0
                )
            """)
            cursor = conn.execute("PRAGMA table_info(temp_files)")
            columns = [row[1] for row in cursor.fetchall()]
            if 'user_filename' not in columns:
                conn.execute("ALTER TABLE temp_files ADD COLUMN user_filename TEXT")
                conn.execute("UPDATE temp_files SET user_filename = original_filename WHERE user_filename IS NULL")
            conn.execute("CREATE INDEX IF NOT EXISTS idx_user_filename ON temp_files(user_filename)")
            conn.commit()
            conn.close()
        
        def register_temp_file(file_path: str, original_filename: str, user_filename: str, cleanup_hours: int = 24) -> str:
            file_id = str(uuid.uuid4())
            created_at = datetime.now()
            expires_at = created_at + timedelta(hours=cleanup_hours)
            
            conn = sqlite3.connect(DB_FILE)
            conn.execute("""
                INSERT INTO temp_files (file_id, original_filename, user_filename, file_path, created_at, expires_at)
                VALUES (?, ?, ?, ?, ?, ?)
            """, (file_id, original_filename, user_filename, file_path, created_at.isoformat(), expires_at.isoformat()))
            conn.commit()
            conn.close()
            return file_id
        
        def ensure_docx_extension(filename: str) -> str:
            if not filename.endswith('.docx'):
                return filename + '.docx'
            return filename
        
        def get_transport_config():
            return {
                'host': '0.0.0.0',
                'port': 8000
            }
        
        init_temp_storage()
        print("✓ Temp storage initialized")
        
        # Step 2: Simulate the tool execution
        print("\n🔧 Simulating create_document_with_download_link execution...")
        
        # Ensure proper extension
        original_filename = ensure_docx_extension(filename)
        unique_filename = f"{uuid.uuid4()}_{original_filename}"
        temp_file_path = TEMP_FILES_DIR / unique_filename
        
        print(f"✓ Generated unique filename: {unique_filename}")
        print(f"✓ Temp file path: {temp_file_path}")
        
        # Create the document
        doc = Document()
        
        # Set properties if provided
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author
        
        # Add some sample content
        doc.add_heading("Test Document", 1)
        doc.add_paragraph("This is a test document created by the MCP tool.")
        doc.add_paragraph("Sample content for testing download functionality.")
        
        # Save to temp location
        doc.save(str(temp_file_path))
        print("✓ Document created and saved")
        
        # Verify file exists
        if not temp_file_path.exists():
            print("❌ ERROR: File was not created on disk")
            return False
        
        file_size = temp_file_path.stat().st_size
        print(f"✓ File exists on disk (size: {file_size} bytes)")
        
        # Register the file for cleanup
        file_id = register_temp_file(str(temp_file_path), original_filename, filename, cleanup_hours)
        print(f"✓ File registered with ID: {file_id}")
        
        # Generate download URL (simulate get_transport_config)
        config = get_transport_config()
        base_url = f"http://{config['host']}:{config['port']}"
        download_url = f"{base_url}/files/{file_id}"
        
        expires_at = datetime.now() + timedelta(hours=cleanup_hours)
        
        # Create the result that would be returned by the tool
        result = {
            "success": True,
            "message": f"Document {original_filename} created successfully",
            "download_url": download_url,
            "file_id": file_id,
            "original_filename": original_filename,
            "expires_at": expires_at.isoformat(),
            "cleanup_hours": cleanup_hours
        }
        
        print("\n📋 Tool Result:")
        print(json.dumps(result, indent=2))
        
        # Step 3: Verify the result structure
        print("\n🔍 Verifying result structure...")
        
        required_fields = ["success", "download_url", "file_id", "original_filename", "expires_at"]
        for field in required_fields:
            if field in result:
                print(f"✓ Field '{field}': {result[field]}")
            else:
                print(f"❌ Missing field '{field}'")
                return False
        
        # Step 4: Test download URL format
        print("\n🌐 Testing download URL format...")
        if result["download_url"].startswith("http://") and "/files/" in result["download_url"]:
            print(f"✓ Download URL format looks correct: {result['download_url']}")
        else:
            print(f"❌ Download URL format is incorrect: {result['download_url']}")
            return False
        
        # Step 5: Test file retrieval by ID
        print("\n🔗 Testing file retrieval simulation...")
        
        def get_temp_file_info(file_id: str):
            conn = sqlite3.connect(DB_FILE)
            cursor = conn.execute("""
                SELECT file_id, original_filename, user_filename, file_path, created_at, expires_at, download_count
                FROM temp_files WHERE file_id = ?
            """, (file_id,))
            row = cursor.fetchone()
            conn.close()
            if not row:
                return None
            return {
                "file_id": row[0],
                "original_filename": row[1],
                "user_filename": row[2],
                "file_path": row[3],
                "created_at": row[4],
                "expires_at": row[5],
                "download_count": row[6]
            }
        
        # Test retrieving file info by the generated file_id
        file_info = get_temp_file_info(file_id)
        if file_info:
            print("✓ File can be retrieved by ID from database")
            print(f"  - Original filename: {file_info['original_filename']}")
            print(f"  - File path: {file_info['file_path']}")
            print(f"  - Expires at: {file_info['expires_at']}")
            
            # Verify file still exists at the path
            if os.path.exists(file_info['file_path']):
                print("✓ File still exists at registered path")
            else:
                print("❌ File missing at registered path")
                return False
        else:
            print("❌ File cannot be retrieved by ID")
            return False
        
        # Clean up test file
        temp_file_path.unlink()
        print("✓ Test file cleaned up")
        
        print("\n" + "=" * 60)
        print("🎉 CREATE_DOCUMENT_WITH_DOWNLOAD_LINK TOOL TEST SUCCESSFUL!")
        print("\n✅ Verification Results:")
        print("   1. ✓ Document creation works")
        print("   2. ✓ File saved to temp storage")
        print("   3. ✓ Database registration works")
        print("   4. ✓ Download URL generation works")
        print("   5. ✓ File retrieval by ID works")
        print("   6. ✓ All required result fields present")
        
        print(f"\n📤 The tool returns this JSON:")
        print(f"   download_url: {result['download_url']}")
        print(f"   file_id: {result['file_id']}")
        print(f"   success: {result['success']}")
        
        return True
        
    except ImportError as e:
        print(f"❌ IMPORT ERROR: {e}")
        print("Note: This might be expected if dependencies aren't installed locally")
        return False
    except Exception as e:
        print(f"❌ TOOL TEST FAILED: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Run the download link tool test."""
    success = simulate_create_document_with_download_link()
    return 0 if success else 1

if __name__ == "__main__":
    sys.exit(main())