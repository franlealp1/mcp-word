"""
Batch document creation tools for ultra-efficient document generation.
Reduces multiple tool calls to single operations for complex documents.
"""

import os
from typing import Dict, List, Optional, Any, Union
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension
from word_document_server.core.styles import ensure_heading_style, ensure_table_style
from word_document_server.utils.document_utils import get_document_properties


async def create_complete_document_with_sections(
    filename: str,
    title: str,
    sections: List[Dict[str, Any]],
    tables: Optional[List[Dict[str, Any]]] = None,
    metadata: Optional[Dict[str, str]] = None,
    cleanup_hours: int = 24
) -> Dict[str, Any]:
    """Create a complete document with multiple sections, tables, and formatting in one call.

    Args:
        filename: Document filename
        title: Main document title
        sections: List of sections with content
        tables: List of tables to insert
        metadata: Document metadata (author, subject, etc.)
        cleanup_hours: Auto-cleanup time

    Section format:
        {
            "heading": "Section Title",
            "level": 1,  # Heading level 1-6
            "content": "Complete paragraph content...",
            "style": "Normal",  # Optional paragraph style
            "table_after": 0,  # Optional: insert table index after this section
            "page_break": False  # Optional: add page break after section
        }

    Table format:
        {
            "rows": 5,
            "cols": 3,
            "data": [["Header1", "Header2", "Header3"], ["Row1Col1", "Row1Col2", "Row1Col3"]],
            "has_header": True,
            "title": "Table Title",
            "style": "Medium Grid 1 Accent 1",
            "alignment": "center"
        }
    """
    try:
        filename = ensure_docx_extension(filename)

        # Check file writeability
        is_writeable, error_message = check_file_writeable(filename)
        if not is_writeable:
            return {
                "success": False,
                "error": f"Cannot create document: {error_message}",
                "sections_processed": 0,
                "tables_created": 0
            }

        # Create new document
        doc = Document()

        # Set metadata if provided
        if metadata:
            if "author" in metadata:
                doc.core_properties.author = metadata["author"]
            if "subject" in metadata:
                doc.core_properties.subject = metadata["subject"]
            if "keywords" in metadata:
                doc.core_properties.keywords = metadata["keywords"]
            if "comments" in metadata:
                doc.core_properties.comments = metadata["comments"]

        # Ensure necessary styles exist
        ensure_heading_style(doc)
        ensure_table_style(doc)

        # Add main title
        if title:
            title_heading = doc.add_heading(title, level=0)
            title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Track statistics
        sections_processed = 0
        tables_created = 0

        # Process each section
        for section_idx, section in enumerate(sections):
            try:
                # Validate section structure
                if not isinstance(section, dict) or "heading" not in section:
                    continue

                heading_text = section.get("heading", "")
                level = max(1, min(6, section.get("level", 1)))  # Ensure level 1-6
                content = section.get("content", "")
                style = section.get("style", "Normal")
                table_after = section.get("table_after")
                page_break = section.get("page_break", False)

                # Add section heading
                if heading_text:
                    doc.add_heading(heading_text, level=level)

                # Add section content (can be multiple paragraphs)
                if content:
                    # Split content by newlines but keep formatting
                    paragraphs = content.split('\n\n')
                    for para_content in paragraphs:
                        if para_content.strip():
                            paragraph = doc.add_paragraph(para_content.strip())
                            try:
                                paragraph.style = style
                            except KeyError:
                                # Style doesn't exist, use Normal
                                paragraph.style = doc.styles['Normal']

                # Insert table after section if specified
                if table_after is not None and tables and 0 <= table_after < len(tables):
                    table_data = tables[table_after]
                    if _insert_table(doc, table_data):
                        tables_created += 1

                # Add page break if requested
                if page_break:
                    doc.add_page_break()

                sections_processed += 1

            except Exception as e:
                # Continue processing other sections even if one fails
                print(f"Error processing section {section_idx}: {e}")
                continue

        # Insert any remaining tables not assigned to sections
        if tables:
            for table_idx, table_data in enumerate(tables):
                # Skip tables already inserted
                table_used = any(s.get("table_after") == table_idx for s in sections if isinstance(s, dict))
                if not table_used:
                    if _insert_table(doc, table_data):
                        tables_created += 1

        # Save document
        doc.save(filename)

        return {
            "success": True,
            "message": f"Complete document '{filename}' created successfully",
            "filename": filename,
            "sections_processed": sections_processed,
            "tables_created": tables_created,
            "total_sections": len(sections) if sections else 0,
            "total_tables": len(tables) if tables else 0
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to create complete document: {str(e)}",
            "filename": filename,
            "sections_processed": 0,
            "tables_created": 0
        }


async def create_complete_document_with_download_link_and_sections(
    filename: str,
    title: str,
    sections: List[Dict[str, Any]],
    tables: Optional[List[Dict[str, Any]]] = None,
    metadata: Optional[Dict[str, str]] = None,
    cleanup_hours: int = 24
) -> Dict[str, Any]:
    """Create a complete document with sections AND return download link in one call.

    This is the ultimate efficiency tool - creates entire document + download link.
    """
    from word_document_server.main import (
        init_temp_storage,
        register_temp_file,
        get_public_base_url,
        TEMP_FILES_DIR
    )
    from pathlib import Path
    import uuid
    from datetime import datetime, timedelta

    try:
        # Ensure temp storage is initialized
        init_temp_storage()

        # Generate unique filename in temp directory
        original_filename = ensure_docx_extension(filename)
        unique_filename = f"{uuid.uuid4()}_{original_filename}"
        temp_file_path = TEMP_FILES_DIR / unique_filename

        # Create document using batch creation
        temp_filename = str(temp_file_path)
        doc_result = await create_complete_document_with_sections(
            temp_filename, title, sections, tables, metadata, cleanup_hours
        )

        if not doc_result.get("success", False):
            return doc_result

        # Register file for download
        file_id = register_temp_file(str(temp_file_path), original_filename, filename, cleanup_hours)

        # Generate download URL
        base_url = get_public_base_url()
        download_url = f"{base_url}/files/{file_id}"
        expires_at = datetime.now() + timedelta(hours=cleanup_hours)

        # Enhanced result with download info
        result = doc_result.copy()
        result.update({
            "download_url": download_url,
            "file_id": file_id,
            "original_filename": original_filename,
            "expires_at": expires_at.isoformat(),
            "cleanup_hours": cleanup_hours,
            "is_temp_file": True
        })

        return result

    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to create document with download link: {str(e)}",
            "filename": filename,
            "download_url": None,
            "sections_processed": 0,
            "tables_created": 0
        }


async def add_multiple_sections_batch(
    filename: str,
    sections: List[Dict[str, Any]]
) -> Dict[str, Any]:
    """Add multiple sections to existing document in one call.

    Args:
        filename: Existing document filename
        sections: List of sections to add (same format as create_complete_document_with_sections)
    """
    try:
        from word_document_server.main import load_document_with_resolver, save_document_with_resolver

        # Load document using resolver
        doc, resolved_path = load_document_with_resolver(filename)

        sections_processed = 0

        # Process each section
        for section_idx, section in enumerate(sections):
            try:
                if not isinstance(section, dict) or "heading" not in section:
                    continue

                heading_text = section.get("heading", "")
                level = max(1, min(6, section.get("level", 1)))
                content = section.get("content", "")
                style = section.get("style", "Normal")
                page_break = section.get("page_break", False)

                # Add section heading
                if heading_text:
                    doc.add_heading(heading_text, level=level)

                # Add section content
                if content:
                    paragraphs = content.split('\n\n')
                    for para_content in paragraphs:
                        if para_content.strip():
                            paragraph = doc.add_paragraph(para_content.strip())
                            try:
                                paragraph.style = style
                            except KeyError:
                                paragraph.style = doc.styles['Normal']

                # Add page break if requested
                if page_break:
                    doc.add_page_break()

                sections_processed += 1

            except Exception as e:
                print(f"Error processing section {section_idx}: {e}")
                continue

        # Save document
        save_document_with_resolver(doc, filename, resolved_path)

        return {
            "success": True,
            "message": f"Added {sections_processed} sections to {filename}",
            "filename": filename,
            "sections_processed": sections_processed,
            "total_sections": len(sections)
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to add sections: {str(e)}",
            "filename": filename,
            "sections_processed": 0
        }


def _insert_table(doc: Document, table_data: Dict[str, Any]) -> bool:
    """Helper function to insert a formatted table into document.

    Args:
        doc: Document object
        table_data: Table configuration dictionary

    Returns:
        bool: True if table was created successfully
    """
    try:
        rows = table_data.get("rows", 0)
        cols = table_data.get("cols", 0)
        data = table_data.get("data", [])
        has_header = table_data.get("has_header", False)
        title = table_data.get("title", "")
        style = table_data.get("style", "Medium Grid 1 Accent 1")
        alignment = table_data.get("alignment", "left")

        if rows <= 0 or cols <= 0:
            return False

        # Add table title if provided
        if title:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.bold = True
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Create table
        table = doc.add_table(rows=rows, cols=cols)

        # Apply table style
        try:
            table.style = style
        except KeyError:
            # Style doesn't exist, use default
            table.style = 'Medium Grid 1 Accent 1'

        # Populate table with data if provided
        if data:
            for row_idx, row_data in enumerate(data):
                if row_idx >= rows:
                    break

                table_row = table.rows[row_idx]
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx >= cols:
                        break

                    cell = table_row.cells[col_idx]
                    cell.text = str(cell_data) if cell_data is not None else ""

                    # Bold header row
                    if has_header and row_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

        # Set table alignment
        if alignment == "center":
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "right":
            table.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        return True

    except Exception as e:
        print(f"Error creating table: {e}")
        return False


async def create_technical_report_template(
    filename: str,
    report_data: Dict[str, Any],
    cleanup_hours: int = 24
) -> Dict[str, Any]:
    """Create a complete technical report using predefined template structure.

    Args:
        filename: Document filename
        report_data: All report data in structured format
        cleanup_hours: Auto-cleanup time

    Report data format:
        {
            "title": "INFORME TÉCNICO INTEGRAL - PRESA ROSARITO",
            "subtitle": "Análisis de Patologías de Hormigón e Informe Hidrológico",
            "metadata": {"author": "...", "subject": "..."},
            "introduction": {
                "content": "Introduction text...",
                "key_data": {"presa": "Rosarito", "rio": "Tiétar"}
            },
            "pathology_report": {
                "general_state": "BUENO EN GENERAL",
                "detected_pathologies": [...],
                "expansion_process": {...},
                "vertical_movements": {"data": [[...]], "headers": [...]}
            },
            "hydrological_report": {
                "basin_characteristics": {...},
                "annual_contributions": {...},
                "flood_studies": {...}
            },
            "conclusions": "Complete conclusions text..."
        }
    """

    try:
        title = report_data.get("title", "INFORME TÉCNICO")
        subtitle = report_data.get("subtitle", "")
        metadata = report_data.get("metadata", {})

        sections = []
        tables = []

        # 1. Introduction section
        if "introduction" in report_data:
            intro_data = report_data["introduction"]
            intro_content = intro_data.get("content", "")

            # Add key data to introduction
            key_data = intro_data.get("key_data", {})
            if key_data:
                intro_content += f"\n\n"
                for key, value in key_data.items():
                    intro_content += f"• {key.title()}: {value}\n"

            sections.append({
                "heading": "1. INTRODUCCIÓN",
                "level": 1,
                "content": intro_content
            })

        # 2. Pathology report section
        if "pathology_report" in report_data:
            path_data = report_data["pathology_report"]

            path_content = f"ESTADO GENERAL: {path_data.get('general_state', 'No especificado')}\n\n"

            # Add detected pathologies
            if "detected_pathologies" in path_data:
                path_content += "PATOLOGÍAS DETECTADAS:\n"
                for pathology in path_data["detected_pathologies"]:
                    path_content += f"• {pathology}\n"
                path_content += "\n"

            # Add expansion process info
            if "expansion_process" in path_data:
                exp_data = path_data["expansion_process"]
                path_content += "PROCESO EXPANSIVO DETECTADO:\n"
                for key, value in exp_data.items():
                    path_content += f"• {key.replace('_', ' ').title()}: {value}\n"

            sections.append({
                "heading": "2. INFORME DE PATOLOGÍAS DE HORMIGÓN",
                "level": 1,
                "content": path_content,
                "table_after": 0
            })

            # Add vertical movements table
            if "vertical_movements" in path_data:
                vm_data = path_data["vertical_movements"]
                tables.append({
                    "rows": len(vm_data.get("data", [])),
                    "cols": len(vm_data.get("headers", [])),
                    "data": [vm_data.get("headers", [])] + vm_data.get("data", []),
                    "has_header": True,
                    "title": "MOVIMIENTOS VERTICALES REGISTRADOS (mm)",
                    "style": "Medium Grid 1 Accent 1"
                })

        # 3. Hydrological report section
        if "hydrological_report" in report_data:
            hydro_data = report_data["hydrological_report"]

            hydro_content = "CARACTERÍSTICAS DE LA CUENCA:\n"

            # Basin characteristics
            if "basin_characteristics" in hydro_data:
                basin_data = hydro_data["basin_characteristics"]
                for key, value in basin_data.items():
                    hydro_content += f"• {key.replace('_', ' ').title()}: {value}\n"
                hydro_content += "\n"

            # Annual contributions
            if "annual_contributions" in hydro_data:
                contrib_data = hydro_data["annual_contributions"]
                hydro_content += "APORTACIONES ANUALES:\n"
                for key, value in contrib_data.items():
                    hydro_content += f"• {key.replace('_', ' ').title()}: {value}\n"
                hydro_content += "\n"

            # Flood studies
            if "flood_studies" in hydro_data:
                flood_data = hydro_data["flood_studies"]
                hydro_content += "ESTUDIOS DE AVENIDA:\n\n"

                for study_name, study_data in flood_data.items():
                    hydro_content += f"{study_name}:\n"
                    for key, value in study_data.items():
                        hydro_content += f"• {key}: {value}\n"
                    hydro_content += "\n"

            sections.append({
                "heading": "3. INFORME HIDROLÓGICO",
                "level": 1,
                "content": hydro_content
            })

        # 4. Conclusions section
        if "conclusions" in report_data:
            sections.append({
                "heading": "4. CONCLUSIONES Y RECOMENDACIONES",
                "level": 1,
                "content": report_data["conclusions"]
            })

        # Create complete document
        result = await create_complete_document_with_download_link_and_sections(
            filename, title, sections, tables, metadata, cleanup_hours
        )

        # Add subtitle if provided
        if subtitle and result.get("success", False):
            # This would require additional implementation to add subtitle
            pass

        # Enhance result with template info
        if result.get("success", False):
            result["template_used"] = "technical_report"
            result["report_sections"] = len(sections)
            result["report_tables"] = len(tables)

        return result

    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to create technical report: {str(e)}",
            "filename": filename,
            "template_used": "technical_report"
        }