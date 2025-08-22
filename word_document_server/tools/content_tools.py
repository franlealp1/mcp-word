"""
Content tools for Word Document Server.

These tools add various types of content to Word documents,
including headings, paragraphs, tables, images, and page breaks.
"""
import os
from typing import List, Optional, Dict, Any
from docx import Document
from docx.shared import Inches, Pt

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension
from word_document_server.utils.document_utils import find_and_replace_text, insert_header_near_text, insert_numbered_list_near_text, insert_line_or_paragraph_near_text, replace_paragraph_block_below_header, replace_block_between_manual_anchors
from word_document_server.core.styles import ensure_heading_style, ensure_table_style


async def add_heading(filename: str, text: str, level: int = 1) -> str:
    """Add a heading to a Word document.
    
    Args:
        filename: Path to the Word document
        text: Heading text
        level: Heading level (1-9, where 1 is the highest level)
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure level is converted to integer
    try:
        level = int(level)
    except (ValueError, TypeError):
        return "Invalid parameter: level must be an integer between 1 and 9"
    
    # Validate level range
    if level < 1 or level > 9:
        return f"Invalid heading level: {level}. Level must be between 1 and 9."
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(filename)
        
        # Ensure heading styles exist
        ensure_heading_style(doc)
        
        # Try to add heading with style
        try:
            heading = doc.add_heading(text, level=level)
            doc.save(filename)
            return f"Heading '{text}' (level {level}) added to {filename}"
        except Exception as style_error:
            # If style-based approach fails, use direct formatting
            paragraph = doc.add_paragraph(text)
            paragraph.style = doc.styles['Normal']
            run = paragraph.runs[0]
            run.bold = True
            # Adjust size based on heading level
            if level == 1:
                run.font.size = Pt(16)
            elif level == 2:
                run.font.size = Pt(14)
            else:
                run.font.size = Pt(12)
            
            doc.save(filename)
            return f"Heading '{text}' added to {filename} with direct formatting (style not available)"
    except Exception as e:
        return f"Failed to add heading: {str(e)}"


async def add_paragraph(filename: str, text: str, style: Optional[str] = None) -> str:
    """Add a paragraph to a Word document.
    
    Args:
        filename: Path to the Word document
        text: Paragraph text
        style: Optional paragraph style name
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(filename)
        paragraph = doc.add_paragraph(text)
        
        if style:
            try:
                paragraph.style = style
            except KeyError:
                # Style doesn't exist, use normal and report it
                paragraph.style = doc.styles['Normal']
                doc.save(filename)
                return f"Style '{style}' not found, paragraph added with default style to {filename}"
        
        doc.save(filename)
        return f"Paragraph added to {filename}"
    except Exception as e:
        return f"Failed to add paragraph: {str(e)}"


async def add_table(filename: str, rows: int, cols: int, data: Optional[List[List[str]]] = None) -> str:
    """Add a table to a Word document.
    
    Args:
        filename: Path to the Word document
        rows: Number of rows in the table
        cols: Number of columns in the table
        data: Optional 2D array of data to fill the table
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(filename)
        table = doc.add_table(rows=rows, cols=cols)
        
        # Try to set the table style
        try:
            table.style = 'Table Grid'
        except KeyError:
            # If style doesn't exist, add basic borders
            pass
        
        # Fill table with data if provided
        if data:
            for i, row_data in enumerate(data):
                if i >= rows:
                    break
                for j, cell_text in enumerate(row_data):
                    if j >= cols:
                        break
                    table.cell(i, j).text = str(cell_text)
        
        doc.save(filename)
        return f"Table ({rows}x{cols}) added to {filename}"
    except Exception as e:
        return f"Failed to add table: {str(e)}"


async def add_picture(filename: str, image_path: str, width: Optional[float] = None) -> str:
    """Add an image to a Word document.
    
    Args:
        filename: Path to the Word document
        image_path: Path to the image file
        width: Optional width in inches (proportional scaling)
    """
    filename = ensure_docx_extension(filename)
    
    # Validate document existence
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Get absolute paths for better diagnostics
    abs_filename = os.path.abspath(filename)
    abs_image_path = os.path.abspath(image_path)
    
    # Validate image existence with improved error message
    if not os.path.exists(abs_image_path):
        return f"Image file not found: {abs_image_path}"
    
    # Check image file size
    try:
        image_size = os.path.getsize(abs_image_path) / 1024  # Size in KB
        if image_size <= 0:
            return f"Image file appears to be empty: {abs_image_path} (0 KB)"
    except Exception as size_error:
        return f"Error checking image file: {str(size_error)}"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(abs_filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(abs_filename)
        # Additional diagnostic info
        diagnostic = f"Attempting to add image ({abs_image_path}, {image_size:.2f} KB) to document ({abs_filename})"
        
        try:
            if width:
                doc.add_picture(abs_image_path, width=Inches(width))
            else:
                doc.add_picture(abs_image_path)
            doc.save(abs_filename)
            return f"Picture {image_path} added to {filename}"
        except Exception as inner_error:
            # More detailed error for the specific operation
            error_type = type(inner_error).__name__
            error_msg = str(inner_error)
            return f"Failed to add picture: {error_type} - {error_msg or 'No error details available'}\nDiagnostic info: {diagnostic}"
    except Exception as outer_error:
        # Fallback error handling
        error_type = type(outer_error).__name__
        error_msg = str(outer_error)
        return f"Document processing error: {error_type} - {error_msg or 'No error details available'}"


async def add_page_break(filename: str) -> str:
    """Add a page break to the document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        doc.add_page_break()
        doc.save(filename)
        return f"Page break added to {filename}."
    except Exception as e:
        return f"Failed to add page break: {str(e)}"


async def add_table_of_contents(filename: str, title: str = "Table of Contents", max_level: int = 3) -> str:
    """Add a table of contents to a Word document based on heading styles.
    
    Args:
        filename: Path to the Word document
        title: Optional title for the table of contents
        max_level: Maximum heading level to include (1-9)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        # Ensure max_level is within valid range
        max_level = max(1, min(max_level, 9))
        
        doc = Document(filename)
        
        # Collect headings and their positions
        headings = []
        for i, paragraph in enumerate(doc.paragraphs):
            # Check if paragraph style is a heading
            if paragraph.style and paragraph.style.name.startswith('Heading '):
                try:
                    # Extract heading level from style name
                    level = int(paragraph.style.name.split(' ')[1])
                    if level <= max_level:
                        headings.append({
                            'level': level,
                            'text': paragraph.text,
                            'position': i
                        })
                except (ValueError, IndexError):
                    # Skip if heading level can't be determined
                    pass
        
        if not headings:
            return f"No headings found in document {filename}. Table of contents not created."
        
        # Create a new document with the TOC
        toc_doc = Document()
        
        # Add title
        if title:
            toc_doc.add_heading(title, level=1)
        
        # Add TOC entries
        for heading in headings:
            # Indent based on level (using tab characters)
            indent = '    ' * (heading['level'] - 1)
            toc_doc.add_paragraph(f"{indent}{heading['text']}")
        
        # Add page break
        toc_doc.add_page_break()
        
        # Get content from original document
        for paragraph in doc.paragraphs:
            p = toc_doc.add_paragraph(paragraph.text)
            # Copy style if possible
            try:
                if paragraph.style:
                    p.style = paragraph.style.name
            except:
                pass
        
        # Copy tables
        for table in doc.tables:
            # Create a new table with the same dimensions
            new_table = toc_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            # Copy cell contents
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        new_table.cell(i, j).text = paragraph.text
        
        # Save the new document with TOC
        toc_doc.save(filename)
        
        return f"Table of contents with {len(headings)} entries added to {filename}"
    except Exception as e:
        return f"Failed to add table of contents: {str(e)}"


async def delete_paragraph(filename: str, paragraph_index: int) -> str:
    """Delete a paragraph from a document.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph to delete (0-based)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate paragraph index
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."
        
        # Delete the paragraph (by removing its content and setting it empty)
        # Note: python-docx doesn't support true paragraph deletion, this is a workaround
        paragraph = doc.paragraphs[paragraph_index]
        p = paragraph._p
        p.getparent().remove(p)
        
        doc.save(filename)
        return f"Paragraph at index {paragraph_index} deleted successfully."
    except Exception as e:
        return f"Failed to delete paragraph: {str(e)}"


async def delete_table(filename: str, table_index: int) -> str:
    """Delete a table from a document.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table to delete (0-based)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate table index
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})."
        
        # Delete the table by removing its element from the document
        table = doc.tables[table_index]
        table._tbl.getparent().remove(table._tbl)
        
        doc.save(filename)
        return f"Table at index {table_index} deleted successfully."
    except Exception as e:
        return f"Failed to delete table: {str(e)}"


async def search_and_replace(filename: str, find_text: str, replace_text: str) -> str:
    """Search for text and replace all occurrences.
    
    Args:
        filename: Path to the Word document
        find_text: Text to search for
        replace_text: Text to replace with
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Perform find and replace
        count = find_and_replace_text(doc, find_text, replace_text)
        
        if count > 0:
            doc.save(filename)
            return f"Replaced {count} occurrence(s) of '{find_text}' with '{replace_text}'."
        else:
            return f"No occurrences of '{find_text}' found."
    except Exception as e:
        return f"Failed to search and replace: {str(e)}"

async def insert_header_near_text_tool(filename: str, target_text: str = None, header_title: str = "", position: str = 'after', header_style: str = 'Heading 1', target_paragraph_index: int = None) -> str:
    """Insert a header (with specified style) before or after the target paragraph. Specify by text or paragraph index."""
    return insert_header_near_text(filename, target_text, header_title, position, header_style, target_paragraph_index)

async def insert_numbered_list_near_text_tool(filename: str, target_text: str = None, list_items: list = None, position: str = 'after', target_paragraph_index: int = None) -> str:
    """Insert a numbered list before or after the target paragraph. Specify by text or paragraph index."""
    return insert_numbered_list_near_text(filename, target_text, list_items, position, target_paragraph_index)

async def insert_line_or_paragraph_near_text_tool(filename: str, target_text: str = None, line_text: str = "", position: str = 'after', line_style: str = None, target_paragraph_index: int = None) -> str:
    """Insert a new line or paragraph (with specified or matched style) before or after the target paragraph. Specify by text or paragraph index."""
    return insert_line_or_paragraph_near_text(filename, target_text, line_text, position, line_style, target_paragraph_index)

async def replace_paragraph_block_below_header_tool(filename: str, header_text: str, new_paragraphs: list, detect_block_end_fn=None) -> str:
    """Reemplaza el bloque de párrafos debajo de un encabezado, evitando modificar TOC."""
    return replace_paragraph_block_below_header(filename, header_text, new_paragraphs, detect_block_end_fn)

async def replace_block_between_manual_anchors_tool(filename: str, start_anchor_text: str, new_paragraphs: list, end_anchor_text: str = None, new_paragraph_style: str = None) -> str:
    """Replace all content between start_anchor_text and end_anchor_text (or next logical header if not provided)."""
    return replace_block_between_manual_anchors(
        doc_path=filename,
        start_anchor=start_anchor_text,
        new_paragraphs=new_paragraphs,
        end_anchor=end_anchor_text,
        new_paragraph_style=new_paragraph_style
    )

async def modify_table_cell(filename: str, table_index: int, row: int, column: int, content: str) -> str:
    """Modify or add content to a specific table cell.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        row: Row index (0-based)
        column: Column index (0-based)
        content: Text content to set in the cell
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
        row = int(row)
        column = int(column)
    except (ValueError, TypeError):
        return "Invalid parameter: table_index, row, and column must be integers"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate table index
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})."
        
        table = doc.tables[table_index]
        
        # Validate row and column indices
        if row < 0 or row >= len(table.rows):
            return f"Invalid row index. Table has {len(table.rows)} rows (0-{len(table.rows)-1})."
        
        if column < 0 or column >= len(table.columns):
            return f"Invalid column index. Table has {len(table.columns)} columns (0-{len(table.columns)-1})."
        
        # Get the target cell
        cell = table.cell(row, column)
        
        # Find a reference cell to copy style from (preferably a non-header cell)
        reference_cell = None
        reference_style = None
        
        # Look for a non-header cell in the same table to copy style from
        for r in range(len(table.rows)):
            for c in range(len(table.columns)):
                if r != row or c != column:  # Don't use the target cell itself
                    try:
                        ref_cell = table.cell(r, c)
                        if ref_cell.text.strip():  # Find a cell with content
                            reference_cell = ref_cell
                            break
                    except IndexError:
                        continue
            if reference_cell:
                break
        
        # If no reference cell found, use the first cell with content
        if not reference_cell:
            for r in range(len(table.rows)):
                for c in range(len(table.columns)):
                    try:
                        ref_cell = table.cell(r, c)
                        if ref_cell.text.strip():
                            reference_cell = ref_cell
                            break
                    except IndexError:
                        continue
                if reference_cell:
                    break
        
        # Clear existing content and add new content
        cell.text = content
        
        # If we found a reference cell, copy its formatting
        if reference_cell and reference_cell.paragraphs:
            ref_paragraph = reference_cell.paragraphs[0]
            if ref_paragraph.runs:
                ref_run = ref_paragraph.runs[0]
                
                # Apply the same formatting to the new content
                if cell.paragraphs:
                    target_paragraph = cell.paragraphs[0]
                    if target_paragraph.runs:
                        target_run = target_paragraph.runs[0]
                        
                        # Copy font properties
                        if hasattr(ref_run.font, 'bold') and ref_run.font.bold is not None:
                            target_run.font.bold = ref_run.font.bold
                        if hasattr(ref_run.font, 'italic') and ref_run.font.italic is not None:
                            target_run.font.italic = ref_run.font.italic
                        if hasattr(ref_run.font, 'underline') and ref_run.font.underline is not None:
                            target_run.font.underline = ref_run.font.underline
                        if hasattr(ref_run.font, 'size') and ref_run.font.size is not None:
                            target_run.font.size = ref_run.font.size
                        if hasattr(ref_run.font, 'name') and ref_run.font.name is not None:
                            target_run.font.name = ref_run.font.name
                        if hasattr(ref_run.font, 'color') and ref_run.font.color.rgb is not None:
                            target_run.font.color.rgb = ref_run.font.color.rgb
        
        doc.save(filename)
        return f"Cell content updated successfully at table {table_index}, row {row}, column {column}."
    except Exception as e:
        return f"Failed to modify table cell: {str(e)}"
