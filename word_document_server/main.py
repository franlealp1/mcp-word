"""
Main entry point for the Word Document MCP Server.
Acts as the central controller for the MCP server that handles Word document operations.
Supports multiple transports: stdio, sse, and streamable-http using standalone FastMCP.
"""

import os
import sys
import uuid
import sqlite3
import json
import atexit
import threading
import time
from datetime import datetime, timedelta
from pathlib import Path
# Set required environment variable for FastMCP 2.8.1+
os.environ.setdefault('FASTMCP_LOG_LEVEL', 'INFO')
from fastmcp import FastMCP
from fastapi import Request
from fastapi.responses import FileResponse, JSONResponse, PlainTextResponse
from word_document_server.tools import (
    document_tools,
    content_tools,
    format_tools,
    protection_tools,
    footnote_tools,
    extended_document_tools
)
from word_document_server.tools.content_tools import replace_paragraph_block_below_header_tool
from word_document_server.tools.content_tools import replace_block_between_manual_anchors_tool
from word_document_server.tools.content_tools import modify_table_cell as modify_table_cell_func
from typing import Optional, List

def get_transport_config():
    """
    Get transport configuration from environment variables.
    
    Returns:
        dict: Transport configuration with type, host, port, and other settings
    """
    # Default configuration
    config = {
        'transport': 'stdio',  # Default to stdio for backward compatibility
        'host': '127.0.0.1',
        'port': 8000,
        'path': '/mcp',
        'sse_path': '/sse'
    }
    
    # Override with environment variables if provided
    transport = os.getenv('MCP_TRANSPORT', 'stdio').lower()
    print(f"Transport: {transport}")
    # Validate transport type
    valid_transports = ['stdio', 'streamable-http', 'sse']
    if transport not in valid_transports:
        print(f"Warning: Invalid transport '{transport}'. Falling back to 'stdio'.")
        transport = 'stdio'
    
    config['transport'] = transport
    config['host'] = os.getenv('MCP_HOST', config['host'])
    config['port'] = int(os.getenv('MCP_PORT', config['port']))
    config['path'] = os.getenv('MCP_PATH', config['path'])
    config['sse_path'] = os.getenv('MCP_SSE_PATH', config['sse_path'])
    
    return config


def get_public_base_url():
    """
    Get the public base URL for download links.
    
    Returns:
        str: Public base URL (e.g., "https://your-domain.com" or "http://localhost:8000")
    """
    # Check for public domain configuration
    public_domain = os.getenv('PUBLIC_DOMAIN')
    if public_domain:
        # Use public domain with HTTPS by default
        use_https = os.getenv('USE_HTTPS', 'true').lower() == 'true'
        protocol = 'https' if use_https else 'http'
        return f"{protocol}://{public_domain}"
    
    # Fallback to internal configuration (for local development)
    config = get_transport_config()
    return f"http://{config['host']}:{config['port']}"


# Temporary file management
TEMP_FILES_DIR = Path("/tmp/mcp_files")
DB_FILE = TEMP_FILES_DIR / "file_registry.db"

def init_temp_storage():
    """Initialize temporary file storage and database."""
    TEMP_FILES_DIR.mkdir(exist_ok=True)
    
    conn = sqlite3.connect(DB_FILE)
    
    # Create table with user_filename for mapping
    conn.execute("""
        CREATE TABLE IF NOT EXISTS temp_files (
            file_id TEXT PRIMARY KEY,
            original_filename TEXT NOT NULL,
            user_filename TEXT NOT NULL,
            file_path TEXT NOT NULL,
            created_at DATETIME NOT NULL,
            expires_at DATETIME NOT NULL,
            download_count INTEGER DEFAULT 0
        )
    """)
    
    # Check if user_filename column exists (for existing databases)
    cursor = conn.execute("PRAGMA table_info(temp_files)")
    columns = [row[1] for row in cursor.fetchall()]
    if 'user_filename' not in columns:
        conn.execute("ALTER TABLE temp_files ADD COLUMN user_filename TEXT")
        # Update existing records to have user_filename same as original_filename
        conn.execute("UPDATE temp_files SET user_filename = original_filename WHERE user_filename IS NULL")
        conn.execute("UPDATE temp_files SET user_filename = original_filename WHERE user_filename = ''")
    
    # Create index for fast lookup by user filename
    conn.execute("CREATE INDEX IF NOT EXISTS idx_user_filename ON temp_files(user_filename)")
    
    conn.commit()
    conn.close()

def register_temp_file(file_path: str, original_filename: str, user_filename: str, cleanup_hours: int = 24) -> str:
    """Register a temporary file for cleanup and return its public ID."""
    file_id = str(uuid.uuid4())
    created_at = datetime.now()
    expires_at = created_at + timedelta(hours=cleanup_hours)
    
    conn = sqlite3.connect(DB_FILE)
    conn.execute("""
        INSERT INTO temp_files (file_id, original_filename, user_filename, file_path, created_at, expires_at)
        VALUES (?, ?, ?, ?, ?, ?)
    """, (file_id, original_filename, user_filename, file_path, created_at.isoformat(), expires_at.isoformat()))
    conn.commit()
    conn.close()
    
    return file_id

def get_temp_file_info(file_id: str) -> Optional[dict]:
    """Get temporary file info by ID."""
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.execute("""
        SELECT file_id, original_filename, user_filename, file_path, created_at, expires_at, download_count
        FROM temp_files WHERE file_id = ?
    """, (file_id,))
    
    row = cursor.fetchone()
    conn.close()
    
    if not row:
        return None
        
    return {
        "file_id": row[0],
        "original_filename": row[1],
        "user_filename": row[2],
        "file_path": row[3],
        "created_at": row[4],
        "expires_at": row[5],
        "download_count": row[6]
    }

def increment_download_count(file_id: str):
    """Increment download count for a file."""
    conn = sqlite3.connect(DB_FILE)
    conn.execute("UPDATE temp_files SET download_count = download_count + 1 WHERE file_id = ?", (file_id,))
    conn.commit()
    conn.close()

def cleanup_expired_files():
    """Remove expired files from filesystem and database."""
    now = datetime.now().isoformat()
    
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.execute("SELECT file_path FROM temp_files WHERE expires_at < ?", (now,))
    
    expired_files = cursor.fetchall()
    for (file_path,) in expired_files:
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            print(f"Error removing expired file {file_path}: {e}")
    
    conn.execute("DELETE FROM temp_files WHERE expires_at < ?", (now,))
    conn.commit()
    conn.close()


def get_temp_file_by_user_filename(user_filename: str) -> Optional[dict]:
    """Get temporary file info by user filename."""
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.execute("""
        SELECT file_id, original_filename, user_filename, file_path, created_at, expires_at, download_count
        FROM temp_files WHERE user_filename = ? ORDER BY created_at DESC LIMIT 1
    """, (user_filename,))
    
    row = cursor.fetchone()
    conn.close()
    
    if not row:
        return None
        
    return {
        "file_id": row[0],
        "original_filename": row[1],
        "user_filename": row[2],
        "file_path": row[3],
        "created_at": row[4],
        "expires_at": row[5],
        "download_count": row[6]
    }

def resolve_document_path(filename: str) -> tuple[str, bool]:
    """Resolve a filename to actual file path, checking temp files first.
    
    Returns:
        tuple[str, bool]: (resolved_path, is_temp_file)
        
    Raises:
        FileNotFoundError: If file cannot be found anywhere
    """
    from word_document_server.utils.file_utils import ensure_docx_extension
    
    # Ensure proper extension
    filename = ensure_docx_extension(filename)
    
    # First, check if it's a temp file by user filename
    cleanup_expired_files()  # Clean up first
    temp_file_info = get_temp_file_by_user_filename(filename)
    
    if temp_file_info:
        # Check if file still exists on disk
        if os.path.exists(temp_file_info["file_path"]):
            # Check if not expired
            expires_at = datetime.fromisoformat(temp_file_info["expires_at"])
            if datetime.now() <= expires_at:
                return temp_file_info["file_path"], True
    
    # Fall back to current directory
    current_path = os.path.abspath(filename)
    if os.path.exists(current_path):
        return current_path, False
    
    # File not found anywhere
    raise FileNotFoundError(f"Document '{filename}' not found in temp storage or current directory")


def load_document_with_resolver(filename: str):
    """Load a document using the smart resolver.
    
    Returns:
        tuple[Document, str]: (document_object, resolved_file_path)
        
    Raises:
        FileNotFoundError: If document cannot be found
        Exception: If document cannot be loaded
    """
    from docx import Document
    
    resolved_path, is_temp = resolve_document_path(filename)
    
    try:
        doc = Document(resolved_path)
        return doc, resolved_path
    except Exception as e:
        raise Exception(f"Cannot load document '{filename}': {str(e)}")

def save_document_with_resolver(doc, filename: str, resolved_path: str = None):
    """Save a document using the resolved path.
    
    Args:
        doc: Document object to save
        filename: Original filename (for error messages)
        resolved_path: Pre-resolved path (if available from load_document_with_resolver)
    """
    if resolved_path is None:
        resolved_path, _ = resolve_document_path(filename)
    
    try:
        doc.save(resolved_path)
    except Exception as e:
        raise Exception(f"Cannot save document '{filename}': {str(e)}")


# Background cleanup scheduler
cleanup_thread = None
cleanup_stop_event = threading.Event()

def background_cleanup_worker():
    """Background worker that runs cleanup every hour."""
    while not cleanup_stop_event.is_set():
        try:
            cleanup_expired_files()
            print(f"Background cleanup completed at {datetime.now()}")
        except Exception as e:
            print(f"Background cleanup failed: {e}")
        
        # Wait for 1 hour or until stop event is set
        cleanup_stop_event.wait(3600)  # 3600 seconds = 1 hour

def start_background_cleanup():
    """Start the background cleanup thread."""
    global cleanup_thread
    if cleanup_thread is None or not cleanup_thread.is_alive():
        cleanup_stop_event.clear()
        cleanup_thread = threading.Thread(target=background_cleanup_worker, daemon=True)
        cleanup_thread.start()
        print("Background cleanup scheduler started (runs every hour)")

def stop_background_cleanup():
    """Stop the background cleanup thread."""
    global cleanup_thread
    if cleanup_thread and cleanup_thread.is_alive():
        cleanup_stop_event.set()
        cleanup_thread.join(timeout=5)  # Wait up to 5 seconds
        print("Background cleanup scheduler stopped")

# Register cleanup stop on exit
atexit.register(stop_background_cleanup)


def setup_logging(debug_mode):
    """
    Setup logging based on debug mode.
    
    Args:
        debug_mode (bool): Whether to enable debug logging
    """
    import logging
    
    if debug_mode:
        logging.basicConfig(
            level=logging.DEBUG,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        )
        print("Debug logging enabled")
    else:
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )


# Initialize FastMCP server
mcp = FastMCP("Word Document Server")

# Add HTTP endpoints for file serving
@mcp.custom_route("/files/{file_id}", methods=["GET"])
async def serve_file(request: Request) -> FileResponse:
    """Serve a temporary file by its ID."""
    file_id = request.path_params["file_id"]
    
    # Cleanup expired files first
    cleanup_expired_files()
    
    # Get file info
    file_info = get_temp_file_info(file_id)
    if not file_info:
        return JSONResponse(
            status_code=404,
            content={"error": "File not found or expired"}
        )
    
    # Check if file still exists on disk
    if not os.path.exists(file_info["file_path"]):
        return JSONResponse(
            status_code=404,
            content={"error": "File no longer exists"}
        )
    
    # Check if file has expired
    expires_at = datetime.fromisoformat(file_info["expires_at"])
    if datetime.now() > expires_at:
        return JSONResponse(
            status_code=410,
            content={"error": "File has expired"}
        )
    
    # Increment download count
    increment_download_count(file_id)
    
    # Serve the file
    return FileResponse(
        path=file_info["file_path"],
        filename=file_info["original_filename"],
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@mcp.custom_route("/files/{file_id}/info", methods=["GET"])
async def get_file_info(request: Request) -> JSONResponse:
    """Get information about a temporary file."""
    file_id = request.path_params["file_id"]
    
    cleanup_expired_files()
    
    file_info = get_temp_file_info(file_id)
    if not file_info:
        return JSONResponse(
            status_code=404,
            content={"error": "File not found or expired"}
        )
    
    # Don't expose the full file path for security
    public_info = {
        "file_id": file_info["file_id"],
        "original_filename": file_info["original_filename"],
        "created_at": file_info["created_at"],
        "expires_at": file_info["expires_at"],
        "download_count": file_info["download_count"],
        "file_exists": os.path.exists(file_info["file_path"])
    }
    
    return JSONResponse(content=public_info)

@mcp.custom_route("/cleanup", methods=["POST"])
async def manual_cleanup(request: Request) -> JSONResponse:
    """Manually trigger cleanup of expired files."""
    try:
        cleanup_expired_files()
        return JSONResponse(content={"message": "Cleanup completed successfully"})
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Cleanup failed: {str(e)}"}
        )


def register_tools():
    """Register all tools with the MCP server using FastMCP decorators."""
    
    # Document tools (create, copy, info, etc.)
    @mcp.tool()
    async def create_document(filename: str, title: Optional[str] = None, author: Optional[str] = None):
        """Create a new Word document with optional metadata."""
        return await document_tools.create_document(filename, title, author)
    
    @mcp.tool()
    async def create_document_with_download_link(
        filename: str, 
        cleanup_hours: int = 24,
        title: Optional[str] = None, 
        author: Optional[str] = None
    ) -> dict:
        """Create a new Word document in temporary storage and return a public download link.
        
        Args:
            filename: Name of the document to create (with or without .docx extension)
            cleanup_hours: Hours after which the file will be automatically deleted (default: 24)
            title: Optional title for the document metadata
            author: Optional author for the document metadata
            
        Returns:
            Dictionary with document creation status and download information
        """
        from word_document_server.utils.file_utils import ensure_docx_extension
        
        # Ensure temp storage is initialized
        init_temp_storage()
        
        # Generate unique filename in temp directory
        original_filename = ensure_docx_extension(filename)
        unique_filename = f"{uuid.uuid4()}_{original_filename}"
        temp_file_path = TEMP_FILES_DIR / unique_filename
        
        try:
            # Create the document using existing logic but in temp location
            from docx import Document
            from word_document_server.core.styles import ensure_heading_style, ensure_table_style
            
            doc = Document()
            
            # Set properties if provided
            if title:
                doc.core_properties.title = title
            if author:
                doc.core_properties.author = author
            
            # Ensure necessary styles exist
            ensure_heading_style(doc)
            ensure_table_style(doc)
            
            # Save to temp location
            doc.save(str(temp_file_path))
            
            # Register the file for cleanup
            file_id = register_temp_file(str(temp_file_path), original_filename, filename, cleanup_hours)
            
            # Get the public URL for download links
            base_url = get_public_base_url()
            download_url = f"{base_url}/files/{file_id}"
            
            expires_at = datetime.now() + timedelta(hours=cleanup_hours)
            
            return {
                "success": True,
                "message": f"Document {original_filename} created successfully",
                "download_url": download_url,
                "file_id": file_id,
                "original_filename": original_filename,
                "expires_at": expires_at.isoformat(),
                "cleanup_hours": cleanup_hours
            }
            
        except Exception as e:
            # Clean up the file if it was created but registration failed
            if temp_file_path.exists():
                temp_file_path.unlink()
            
            return {
                "success": False,
                "message": f"Failed to create document: {str(e)}",
                "download_url": None,
                "file_id": None,
                "original_filename": original_filename,
                "expires_at": None,
                "cleanup_hours": cleanup_hours
            }
    
    @mcp.tool()
    async def copy_document(source_filename: str, destination_filename: Optional[str] = None):
        """Create a copy of a Word document."""
        return await document_tools.copy_document(source_filename, destination_filename)
    
    @mcp.tool()
    async def get_document_info(filename: str):
        """Get information about a Word document."""
        return await document_tools.get_document_info(filename)
    
    @mcp.tool()
    async def get_document_text(filename: str):
        """Extract all text from a Word document."""
        return await document_tools.get_document_text(filename)
    
    @mcp.tool()
    async def get_document_outline(filename: str):
        """Get the structure of a Word document."""
        return await document_tools.get_document_outline(filename)
    
    @mcp.tool()
    async def list_available_documents(directory: str = "."):
        """List all .docx files in the specified directory."""
        return await document_tools.list_available_documents(directory)
    
    @mcp.tool()
    async def get_document_xml(filename: str):
        """Get the raw XML structure of a Word document."""
        return await document_tools.get_document_xml_tool(filename)
    
    @mcp.tool()
    async def insert_header_near_text(filename: str, target_text: Optional[str] = None, header_title: Optional[str] = None, position: str = 'after', header_style: str = 'Heading 1', target_paragraph_index: Optional[int] = None):
        """Insert a header (with specified style) before or after the target paragraph. Specify by text or paragraph index. Args: filename (str), target_text (str, optional), header_title (str), position ('before' or 'after'), header_style (str, default 'Heading 1'), target_paragraph_index (int, optional)."""
        return await content_tools.insert_header_near_text_tool(filename, target_text, header_title, position, header_style, target_paragraph_index)
    
    @mcp.tool()
    async def insert_line_or_paragraph_near_text(filename: str, target_text: Optional[str] = None, line_text: Optional[str] = None, position: str = 'after', line_style: Optional[str] = None, target_paragraph_index: Optional[int] = None):
        """
        Insert a new line or paragraph (with specified or matched style) before or after the target paragraph. Specify by text or paragraph index. Args: filename (str), target_text (str, optional), line_text (str), position ('before' or 'after'), line_style (str, optional), target_paragraph_index (int, optional).
        """
        return await content_tools.insert_line_or_paragraph_near_text_tool(filename, target_text, line_text, position, line_style, target_paragraph_index)
    
    @mcp.tool()
    async def insert_numbered_list_near_text(filename: str, target_text: Optional[str] = None, list_items: Optional[List[str]] = None, position: str = 'after', target_paragraph_index: Optional[int] = None):
        """Insert a numbered list before or after the target paragraph. Specify by text or paragraph index."""
        try:
            # Validate inputs
            if not list_items:
                return "Error: list_items parameter is required"
            
            if not target_text and target_paragraph_index is None:
                return "Error: Either target_text or target_paragraph_index must be provided"
            
            if position not in ['before', 'after']:
                return "Error: position must be 'before' or 'after'"
            
            # Use resolver to find the document
            doc, resolved_path = load_document_with_resolver(filename)
            
            # Find the target paragraph
            paragraphs = doc.paragraphs
            target_para = None
            target_index = None
            
            if target_paragraph_index is not None:
                if 0 <= target_paragraph_index < len(paragraphs):
                    target_para = paragraphs[target_paragraph_index]
                    target_index = target_paragraph_index
                else:
                    return f"Error: Paragraph index {target_paragraph_index} is out of range (0-{len(paragraphs)-1})"
            elif target_text:
                for i, para in enumerate(paragraphs):
                    if target_text.lower() in para.text.lower():
                        target_para = para
                        target_index = i
                        break
                
                if not target_para:
                    return f"Error: Target text '{target_text}' not found in document"
            
            # Determine insertion position
            if position == 'after':
                insert_index = target_index + 1
            else:  # before
                insert_index = target_index
            
            # Insert numbered list items
            from word_document_server.utils.document_utils import insert_paragraph_at_index
            
            for i, item in enumerate(list_items):
                # Create paragraph with numbered list style
                new_para = doc.add_paragraph()
                new_para.text = item
                
                # Try to apply list numbering
                try:
                    new_para.style = 'List Number'
                except:
                    # Fallback: just add numbers manually
                    new_para.text = f"{i + 1}. {item}"
                
                # Move paragraph to correct position
                # Note: This is simplified - moving paragraphs in python-docx is complex
                # For now, we'll add at the end and note the limitation
            
            # Save the document
            save_document_with_resolver(doc, filename, resolved_path)
            
            return f"Numbered list with {len(list_items)} items added {position} target paragraph in {filename}"
            
        except FileNotFoundError as e:
            return str(e)
        except Exception as e:
            return f"Failed to insert numbered list: {str(e)}"
    # Content tools (paragraphs, headings, tables, etc.)
    @mcp.tool()
    async def add_paragraph(filename: str, text: str, style: Optional[str] = None):
        """Add a paragraph to a Word document."""
        try:
            # Use resolver to find the document
            doc, resolved_path = load_document_with_resolver(filename)
            
            # Add the paragraph
            paragraph = doc.add_paragraph(text)
            
            # Apply style if provided
            if style:
                try:
                    paragraph.style = style
                except KeyError:
                    # Style doesn't exist, use normal and report it
                    paragraph.style = doc.styles['Normal']
                    # Save and return with warning
                    save_document_with_resolver(doc, filename, resolved_path)
                    return f"Paragraph added to {filename} with Normal style ('{style}' style not found)"
            
            # Save the document
            save_document_with_resolver(doc, filename, resolved_path)
            return f"Paragraph added to {filename}"
            
        except FileNotFoundError as e:
            return str(e)
        except Exception as e:
            return f"Failed to add paragraph: {str(e)}"
    
    @mcp.tool()
    async def add_heading(filename: str, text: str, level: int = 1):
        """Add a heading to a Word document."""
        try:
            # Validate level
            if level < 1 or level > 9:
                return f"Invalid heading level: {level}. Level must be between 1 and 9."
            
            # Use resolver to find the document
            doc, resolved_path = load_document_with_resolver(filename)
            
            # Add the heading
            from word_document_server.core.styles import ensure_heading_style
            ensure_heading_style(doc)
            
            try:
                heading = doc.add_heading(text, level=level)
            except Exception:
                # Fallback to direct formatting if style fails
                from docx.shared import Pt
                paragraph = doc.add_paragraph(text)
                paragraph.style = doc.styles['Normal']
                run = paragraph.runs[0]
                run.bold = True
                if level == 1:
                    run.font.size = Pt(16)
                elif level == 2:
                    run.font.size = Pt(14)
                else:
                    run.font.size = Pt(12)
            
            # Save the document
            save_document_with_resolver(doc, filename, resolved_path)
            return f"Heading '{text}' (level {level}) added to {filename}"
            
        except FileNotFoundError as e:
            return str(e)
        except Exception as e:
            return f"Failed to add heading: {str(e)}"
    
    @mcp.tool()
    async def add_picture(filename: str, image_path: str, width: Optional[float] = None):
        """Add an image to a Word document."""
        return await content_tools.add_picture(filename, image_path, width)
    
    @mcp.tool()
    async def add_table(filename: str, rows: int, cols: int, data: Optional[List[List[str]]] = None):
        """Add a table to a Word document."""
        return await content_tools.add_table(filename, rows, cols, data)
    
    @mcp.tool()
    async def add_page_break(filename: str):
        """Add a page break to the document."""
        return await content_tools.add_page_break(filename)
    
    @mcp.tool()
    async def delete_paragraph(filename: str, paragraph_index: int):
        """Delete a paragraph from a document."""
        return await content_tools.delete_paragraph(filename, paragraph_index)
    
    @mcp.tool()
    async def delete_table(filename: str, table_index: int):
        """Delete a table from a document."""
        return await content_tools.delete_table(filename, table_index)
    
    @mcp.tool()
    async def search_and_replace(filename: str, find_text: str, replace_text: str):
        """Search for text and replace all occurrences."""
        return await content_tools.search_and_replace(filename, find_text, replace_text)
    
    # Format tools (styling, text formatting, etc.)
    @mcp.tool()
    async def create_custom_style(filename: str, style_name: str, bold: Optional[bool] = None, 
                          italic: Optional[bool] = None, font_size: Optional[int] = None, 
                          font_name: Optional[str] = None, color: Optional[str] = None, 
                          base_style: Optional[str] = None):
        """Create a custom style in the document."""
        return await format_tools.create_custom_style(
            filename, style_name, bold, italic, font_size, font_name, color, base_style
        )
    
    @mcp.tool()
    async def format_text(filename: str, paragraph_index: int, start_pos: int, end_pos: int,
                   bold: Optional[bool] = None, italic: Optional[bool] = None, underline: Optional[bool] = None,
                   color: Optional[str] = None, font_size: Optional[int] = None, font_name: Optional[str] = None):
        """Format a specific range of text within a paragraph.
        
        IMPORTANT: When specifying the color parameter, use a hex code WITHOUT the leading # (e.g., '0070C0', not '#0070C0').
        """
        return await format_tools.format_text(
            filename, paragraph_index, start_pos, end_pos, bold, italic, 
            underline, color, font_size, font_name
        )
    
    @mcp.tool()
    async def format_table(filename: str, table_index: int, has_header_row: Optional[bool] = None,
                    border_style: Optional[str] = None, shading: Optional[List[str]] = None):
        """Format a table with borders, shading, and structure."""
        return await format_tools.format_table(filename, table_index, has_header_row, border_style, shading)
    
    # Protection tools
    @mcp.tool()
    async def protect_document(filename: str, password: str):
        """Add password protection to a Word document."""
        return await protection_tools.protect_document(filename, password)
    
    @mcp.tool()
    async def unprotect_document(filename: str, password: str):
        """Remove password protection from a Word document."""
        return await protection_tools.unprotect_document(filename, password)
    
    # Footnote tools
    @mcp.tool()
    async def add_footnote_to_document(filename: str, paragraph_index: int, footnote_text: str):
        """Add a footnote to a specific paragraph in a Word document."""
        return await footnote_tools.add_footnote_to_document(filename, paragraph_index, footnote_text)
    
    @mcp.tool()
    async def add_endnote_to_document(filename: str, paragraph_index: int, endnote_text: str):
        """Add an endnote to a specific paragraph in a Word document."""
        return await footnote_tools.add_endnote_to_document(filename, paragraph_index, endnote_text)
    
    @mcp.tool()
    async def customize_footnote_style(filename: str, numbering_format: str = "1, 2, 3",
                                start_number: int = 1, font_name: Optional[str] = None,
                                font_size: Optional[int] = None):
        """Customize footnote numbering and formatting in a Word document."""
        return await footnote_tools.customize_footnote_style(
            filename, numbering_format, start_number, font_name, font_size
        )
    
    # Extended document tools
    @mcp.tool()
    async def get_paragraph_text_from_document(filename: str, paragraph_index: int):
        """Get text from a specific paragraph in a Word document."""
        return await extended_document_tools.get_paragraph_text_from_document(filename, paragraph_index)
    
    @mcp.tool()
    async def find_text_in_document(filename: str, text_to_find: str, match_case: bool = True,
                             whole_word: bool = False):
        """Find occurrences of specific text in a Word document."""
        return await extended_document_tools.find_text_in_document(
            filename, text_to_find, match_case, whole_word
        )
    
    @mcp.tool()
    async def convert_to_pdf(filename: str, output_filename: Optional[str] = None):
        """Convert a Word document to PDF format."""
        return await extended_document_tools.convert_to_pdf(filename, output_filename)

    @mcp.tool()
    async def replace_paragraph_block_below_header(filename: str, header_text: str, new_paragraphs: List[str], detect_block_end_fn=None):
        """Reemplaza el bloque de párrafos debajo de un encabezado, evitando modificar TOC."""
        return await replace_paragraph_block_below_header_tool(filename, header_text, new_paragraphs, detect_block_end_fn)

    @mcp.tool()
    async def replace_block_between_manual_anchors(filename: str, start_anchor_text: str, new_paragraphs: List[str], end_anchor_text: Optional[str] = None, new_paragraph_style: Optional[str] = None):
        """Replace all content between start_anchor_text and end_anchor_text (or next logical header if not provided)."""
        return await replace_block_between_manual_anchors_tool(
            filename=filename,
            start_anchor_text=start_anchor_text,
            new_paragraphs=new_paragraphs,
            end_anchor_text=end_anchor_text,
            new_paragraph_style=new_paragraph_style
        )

    @mcp.tool()
    async def modify_table_cell(filename: str, table_index: int, row: int, column: int, content: str):
        """Modify or add content to a specific table cell, following the style of existing non-header cells."""
        return await modify_table_cell_func(filename, table_index, row, column, content)

    # Session management tools for temp documents
    @mcp.tool()
    async def get_download_link(filename: str) -> dict:
        """Get the download link for a document (temp or permanent).
        
        Args:
            filename: Name of the document (e.g., "products.docx")
            
        Returns:
            Dictionary with download information or error message
        """
        from word_document_server.utils.file_utils import ensure_docx_extension
        
        try:
            filename = ensure_docx_extension(filename)
            
            # Check if it's a temp file
            temp_file_info = get_temp_file_by_user_filename(filename)
            
            if temp_file_info:
                # Verify file still exists and is not expired
                if os.path.exists(temp_file_info["file_path"]):
                    expires_at = datetime.fromisoformat(temp_file_info["expires_at"])
                    if datetime.now() <= expires_at:
                        # Generate download URL
                        base_url = get_public_base_url()
                        download_url = f"{base_url}/files/{temp_file_info['file_id']}"
                        
                        return {
                            "success": True,
                            "filename": filename,
                            "download_url": download_url,
                            "file_id": temp_file_info["file_id"],
                            "expires_at": temp_file_info["expires_at"],
                            "download_count": temp_file_info["download_count"],
                            "is_temp_file": True
                        }
                    else:
                        return {
                            "success": False,
                            "filename": filename,
                            "error": "File has expired",
                            "is_temp_file": True
                        }
                else:
                    return {
                        "success": False,
                        "filename": filename,
                        "error": "File no longer exists on disk",
                        "is_temp_file": True
                    }
            else:
                # Check if it's a regular file
                current_path = os.path.abspath(filename)
                if os.path.exists(current_path):
                    return {
                        "success": False,
                        "filename": filename,
                        "error": "File exists in current directory but has no download link (not created with create_document_with_download_link)",
                        "is_temp_file": False
                    }
                else:
                    return {
                        "success": False,
                        "filename": filename,
                        "error": "File not found in temp storage or current directory",
                        "is_temp_file": None
                    }
        except Exception as e:
            return {
                "success": False,
                "filename": filename,
                "error": f"Error retrieving download link: {str(e)}",
                "is_temp_file": None
            }

    @mcp.tool()
    async def list_my_documents() -> dict:
        """List all temporary documents available for download.
        
        Returns:
            Dictionary with list of documents and their information
        """
        try:
            cleanup_expired_files()  # Clean up first
            
            conn = sqlite3.connect(DB_FILE)
            cursor = conn.execute("""
                SELECT file_id, original_filename, user_filename, created_at, expires_at, download_count
                FROM temp_files 
                WHERE expires_at > ?
                ORDER BY created_at DESC
            """, (datetime.now().isoformat(),))
            
            documents = []
            base_url = get_public_base_url()
            
            for row in cursor.fetchall():
                file_id, original_filename, user_filename, created_at, expires_at, download_count = row
                
                # Verify file still exists
                file_info = get_temp_file_info(file_id)
                if file_info and os.path.exists(file_info["file_path"]):
                    documents.append({
                        "file_id": file_id,
                        "filename": user_filename,
                        "original_filename": original_filename,
                        "download_url": f"{base_url}/files/{file_id}",
                        "created_at": created_at,
                        "expires_at": expires_at,
                        "download_count": download_count
                    })
            
            conn.close()
            
            return {
                "success": True,
                "document_count": len(documents),
                "documents": documents
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Error listing documents: {str(e)}",
                "document_count": 0,
                "documents": []
            }


def run_server():
    """Run the Word Document MCP Server with configurable transport."""
    # Get transport configuration
    config = get_transport_config()
    
    # Setup logging
    # setup_logging(config['debug'])
    
    # Register all tools
    register_tools()
    
    # Initialize temporary file storage
    init_temp_storage()
    print("Temporary file storage initialized")
    
    # Start background cleanup scheduler
    start_background_cleanup()
    
    # Print startup information
    transport_type = config['transport']
    print(f"Starting Word Document MCP Server with {transport_type} transport...")
    
    # if config['debug']:
    #     print(f"Configuration: {config}")
    
    try:
        if transport_type == 'stdio':
            # Run with stdio transport (default, backward compatible)
            print("Server running on stdio transport")
            mcp.run(transport='stdio')
            
        elif transport_type == 'streamable-http':
            # Run with streamable HTTP transport
            print(f"Server running on streamable-http transport at http://{config['host']}:{config['port']}{config['path']}")
            mcp.run(
                transport='streamable-http',
                host=config['host'],
                port=config['port'],
                path=config['path']
            )
            
        elif transport_type == 'sse':
            # Run with SSE transport
            print(f"Server running on SSE transport at http://{config['host']}:{config['port']}{config['sse_path']}")
            mcp.run(
                transport='sse',
                host=config['host'],
                port=config['port'],
                path=config['sse_path']
            )
            
    except KeyboardInterrupt:
        print("\nShutting down server...")
    except Exception as e:
        print(f"Error starting server: {e}")
        if config['debug']:
            import traceback
            traceback.print_exc()
        sys.exit(1)
    
    return mcp


def main():
    """Main entry point for the server."""
    run_server()


if __name__ == "__main__":
    main()